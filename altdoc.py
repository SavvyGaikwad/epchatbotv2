from docx import Document
import os
import json
import uuid
import shutil
from typing import List, Dict, Union
from pathlib import Path

def docx_to_structured_json(docx_path: str, 
                          output_json_path: str, 
                          image_output_dir: str = 'extracted_images',
                          github_base_url: str = None) -> Dict:
    """
    Convert Word document to structured JSON with proper image handling.
    
    Args:
        docx_path: Path to the input Word document
        output_json_path: Path to save the output JSON
        image_output_dir: Directory to save extracted images
        github_base_url: Optional base URL if images will be hosted on GitHub
    
    Returns:
        Dictionary containing the structured content
    """
    
    # Create output directory if it doesn't exist
    Path(image_output_dir).mkdir(parents=True, exist_ok=True)
    
    document = Document(docx_path)
    result = {
        "document_title": os.path.splitext(os.path.basename(docx_path))[0],
        "content": []
    }
    
    current_section = None
    image_counter = 1
    
    for paragraph in document.paragraphs:
        # Check if this is a heading (section title)
        if paragraph.style.name.startswith('Heading'):
            # Save previous section if exists
            if current_section:
                result["content"].append(current_section)
            
            # Start new section
            current_section = {
                "type": "section",
                "title": paragraph.text.strip(),
                "content": []
            }
        else:
            # Only process if we're in a section
            if current_section is None:
                continue
                
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Determine content type based on formatting or text patterns
            content_type = "info"
            if text.startswith(('• ', '- ')) or any(text.lstrip().startswith(f"{i}.") for i in range(1, 20)):
                content_type = "info"
            elif text.split('.')[0].isdigit():  # Step if starts with number
                content_type = "step"
                
            # Add text content
            current_section["content"].append({
                "type": content_type,
                "text": text
            })
    
    # Process images and add them to the appropriate locations
    rels = document.part.rels
    for rel in rels.values():
        if "image" not in rel.reltype:
            continue
            
        image_part = rel.target_part
        image_data = image_part.blob
        
        # Generate unique filename
        ext = image_part.content_type.split('/')[-1]
        filename = f"image_{image_counter}.{ext}"
        filepath = os.path.join(image_output_dir, filename)
        
        # Save image to file
        with open(filepath, 'wb') as f:
            f.write(image_data)
        
        # Determine image path (local or GitHub)
        image_path = filepath
        if github_base_url:
            image_path = f"{github_base_url}/{filename}"
        
        # Find the paragraph that contains this image
        for paragraph in document.paragraphs:
            if '<w:drawing>' in paragraph._p.xml and str(rel.rId) in paragraph._p.xml:
                # Add image to content
                if current_section:
                    current_section["content"].append({
                        "type": "media",
                        "path": image_path
                    })
                image_counter += 1
                break
    
    # Add the last section if exists
    if current_section:
        result["content"].append(current_section)
    
    # Save to JSON
    with open(output_json_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, indent=2, ensure_ascii=False)
    
    return result

# Example usage with GitHub hosting
if __name__ == "__main__":
    docx_file = r"C:\Users\hp\Downloads\English-20250601T110944Z-1-001\user manual\Admin Settings\Administrator  Settings.docx" # Your Word document
    json_output = r"C:\Users\hp\OneDrive\Desktop\epchatbot-finalvr-main\data\Administrator Settings.json"  # Output JSON file
    image_dir = "extracted_images"  # Local image directory
    github_url = "https://github.com/yourusername/yourrepo/blob/main/images"  # Your GitHub URL
    
    # Convert with GitHub hosting
    extracted_data = docx_to_structured_json(
        docx_file,
        json_output,
        image_output_dir=image_dir,
        github_base_url=github_url
    )
    
    print(f"Successfully converted {docx_file} to {json_output}")
    print(f"Images saved to {image_dir}")